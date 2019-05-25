"""
Make 'span' in tags dict a stack
maybe do the same for all tags in case of unclosed tags?

the idea is that there is a method that converts html files into docx
but also have api methods that let user have more control e.g. so they
can nest calls to something like 'convert_chunk' in loops

also, user can still access Document methods directly (risky?)

deal with tables
ignore 'code'
"""
from html.parser import HTMLParser
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR, WD_ALIGN_PARAGRAPH

import re

style_map = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    's': 'strikethrough',
    'u': 'underline',
    'p': 'paragraph',
    'ul': 'unordered list',
    'li': 'list item',
    # TODO sort this out?
}

style_type = {
    'b': 'font',
    'strong': 'font',
}

last_tag = None

class MyHTMLParser(HTMLParser):
    
    # def __init__(self):
    #     self.styles = {}

    def handle_starttag(self, tag, attrs):
        # print("<start tag:", tag)
        # if attrs:
        #     print(attrs)
        self.tags[tag] = dict(attrs)
        self.stack.append((tag, dict(attrs)))
        if tag == 'p':
            self.paragraph = self.doc.add_paragraph()
            if attrs:
                d = dict(attrs)
                if 'style' in d:
                    style = d['style'].replace(" ", '').split(';')
                    style = dict([x.split(':') for x in style if ':' in x])
                    print(style)
                    if 'text-align' in style:
                        align = style['text-align']
                        if align == 'center':
                            self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif align == 'right':
                            self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        elif align == 'justify':
                            self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    if 'margin-left' in style:
                        margin = style['margin-left']
                        margin = int(re.sub(r'[a-z]+', '', margin))
                        # if margin is in px
                        self.paragraph.paragraph_format.left_indent = Inches(min(margin // 10 * 0.25, 5.5))
                        
        elif tag == 'li':
            if 'ol' in self.tags:
                self.paragraph = self.doc.add_paragraph(style='List Number')
            elif 'ul' in self.tags:
                self.paragraph = self.doc.add_paragraph(style='List Bullet')
            else:
                self.paragraph = self.doc.add_paragraph(style='List Bullet')
        elif tag[0] == 'h' and len(tag) == 2:
            try:
                h_size = int(tag[1])
                self.paragraph = self.doc.add_heading(level=h_size)
            except:
                import pdb; pdb.set_trace()

    def handle_endtag(self, tag):
        # print("end tag>:", tag)
        try:
            if tag in self.tags:
                self.tags.pop(tag)
            # check tag is top of stack?
            if tag == self.stack[-1][0]:
                self.stack.pop()
        except:
            import pdb; pdb.set_trace()

    def handle_data(self, data):
        # print("DATA:", data)
        print(self.stack)
        # check through tags and determine data type and styles to add
        
        # check for text type (paragraph, list etc)
        run = self.paragraph.add_run(data)
        for item in self.stack:
            if item[0] not in ['p', 'li']:
                tag = item[0]
                if 'style' in item[1]:
                    style = item[1]['style'].replace(' ', '').split(';')
                    style = dict([x.split(':') for x in style if ':' in x])
                    print(style)
                    if tag == 'span':
                        if 'color' in style:
                            color = re.sub(r'[a-z()]+', '', style['color'])
                            colors = [int(x) for x in color.split(',')]
                            run.font.color.rgb = RGBColor(*colors)
                            print(colors, run.font.color.rgb, run)
                        if 'background-color' in style:
                            color = color = re.sub(r'[a-z()]+', '', style['background-color'])
                            colors = [int(x) for x in color.split(',')]
                            run.font.highlight_color = WD_COLOR.BLUE # some kind of color mapping with thresholds?
                            # use subset of all colors and find closest with euclidean distance
                    
                else: # single element tag like bold/italic
                    if tag == 'strong':
                        run.font.bold = True
                    elif tag == 'em':
                        run.font.italic = True
                    elif tag == 'u':
                        run.font.underline = True
                    elif tag == 's':
                        run.font.strike = True

                


    def set_initial_attrs(self):
        self.tags = {}
        self.stack = []
        self.doc = Document()
        self.last_tag = None

    def run(self, html, filename=None):
        self.set_initial_attrs()
        self.feed(html)
        print(self.stack)
        if filename:
            self.doc.save('f1.docx')
        else:
            self.doc.save('%s.docx' % filename)


parser = MyHTMLParser()
# parser.feed('<html><head><title style="color: blue">Test</title></head>'
#             '<body><h1>Parse me!</h1></body></html>')
parser.set_initial_attrs()
parser.feed("""<p>My line <span width="10" style="color: rgb(235, 107, 86);">goes he</span>re</p><p><span style="background-color: rgb(251, 160, 38);">Background color</span></p><p><span style="background-color: rgb(71, 85, 119); color: rgb(255, 255, 255);">This sentence has background and&nbsp;</span><span style="background-color: rgb(71, 85, 119);"><span style="color: rgb(247, 218, 100);">text color</span><span style="color: rgb(255, 255, 255);">&nbsp;</span></span><span style="background-color: rgb(71, 85, 119); color: rgb(255, 255, 255);">and<strong> bold </strong><em>italic</em> <u>underlined</u> <s>strike</s> styles</span></p><p>List</p><pre>2 + 3 = 5↵this is code</pre><h1>heading 1</h1><ol><li>A list for reals</li><li>Second item</li></ol><ul style="list-style-type: circle;"><li>Unorderd list</li><li>with circle markers</li></ul><p>Align left Align leftAlign leftAlign leftAlign leftAlign leftAlign leftAlign leftAlign leftAlign left</p><p style="text-align: center;">Align center Align center Align center Align center Align center Align center Align center Align center Align center</p><p style="text-align: right;">Align Right. Align Right. Align Right. Align Right. Align Right. Align Right. Align Right. Align Right.</p><p style="text-align: justify;">This sentence is justified. This sentence is justified. This sentence is justified. This sentence is justified. This sentence is justified.</p><p style="text-align: justify;">Indent 0</p><p style="text-align: justify; margin-left: 20px;">Indent 1</p><p style="text-align: justify; margin-left: 40px;">Indent 2</p><p style="text-align: justify; margin-left: 60px;">Indent 3</p><p style="text-align: justify; margin-left: 80px;">Indent 4</p><p style="text-align: justify; margin-left: 580px;">Indent max?</p><p style="text-align: left;">asdfsa</p><p style="text-align: left;"><a class="fr-green fr-strong" href="https://github.com" rel="noopener noreferrer" target="_blank">link</a></p>""")
parser.doc.save('t1.docx')