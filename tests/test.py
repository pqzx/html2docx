import os
from pathlib import Path
import unittest
from docx import Document
from .context import HtmlToDocx, test_dir


class OutputTest(unittest.TestCase):

    @staticmethod
    def get_html_from_file(filename: str):
        file_path = Path(test_dir) / Path(filename)
        with open(file_path, 'r') as f:
            html = f.read()
        return html

    @classmethod
    def setUpClass(cls):
        cls.document = Document()
        cls.text1 = cls.get_html_from_file('text1.html')
        cls.table_html = cls.get_html_from_file('tables1.html')
        cls.table2_html = cls.get_html_from_file('tables2.html')

    @classmethod
    def tearDownClass(cls):
        outputpath = os.path.join(test_dir, 'test.docx')
        cls.document.save(outputpath)

    def setUp(self):
        self.parser = HtmlToDocx()

    def test_html_with_images_links_style(self):
        self.document.add_heading(
            'Test: add regular html with images, links and some formatting to document',
            level=1
        )
        self.parser.add_html_to_document(self.text1, self.document)

    def test_html_with_default_paragraph_style(self):
        self.document.add_heading(
            'Test: add regular html with a default paragraph style defined',
            level=1
        )
        self.parser.paragraph_style = 'Quote'
        self.parser.add_html_to_document(self.text1, self.document)

    def test_add_html_to_table_cell_with_default_paragraph_style(self):
        self.document.add_heading(
            'Test: regular html to table cell with a default paragraph style defined',
            level=1
        )
        self.parser.paragraph_style = 'Quote'
        table = self.document.add_table(2, 2, style='Table Grid')
        cell = table.cell(1, 1)
        self.parser.add_html_to_document(self.text1, cell)

    def test_add_html_to_table_cell(self):
        self.document.add_heading(
            'Test: regular html with images, links, some formatting to table cell',
            level=1
        )
        table = self.document.add_table(2,2, style='Table Grid')
        cell = table.cell(1,1)
        self.parser.add_html_to_document(self.text1, cell)

    def test_add_html_skip_images(self):
        self.document.add_heading(
            'Test: regular html with images, but skip adding images',
            level=1
        )
        self.parser.options['images'] = False
        self.parser.add_html_to_document(self.text1, self.document)

    def test_add_html_with_tables(self):
        self.document.add_heading(
            'Test: add html with tables',
            level=1
        )
        self.parser.add_html_to_document(self.table_html, self.document)

    def test_add_html_with_tables_accent_style(self):
        self.document.add_heading(
            'Test: add html with tables with accent',
        )
        self.parser.table_style = 'Light Grid Accent 6'
        self.parser.add_html_to_document(self.table_html, self.document)

    def test_add_html_with_tables_basic_style(self):
        self.document.add_heading(
            'Test: add html with tables with basic style',
        )
        self.parser.table_style = 'TableGrid'
        self.parser.add_html_to_document(self.table_html, self.document)

    def test_add_nested_tables(self):
        self.document.add_heading(
            'Test: add nested tables',
        )
        self.parser.add_html_to_document(self.table2_html, self.document)

    def test_add_nested_tables_basic_style(self):
        self.document.add_heading(
            'Test: add nested tables with basic style',
        )
        self.parser.table_style = 'TableGrid'
        self.parser.add_html_to_document(self.table2_html, self.document)

    def test_add_nested_tables_accent_style(self):
        self.document.add_heading(
            'Test: add nested tables with accent style',
        )
        self.parser.table_style = 'Light Grid Accent 6'
        self.parser.add_html_to_document(self.table2_html, self.document)

    def test_add_html_skip_tables(self):
        # broken until feature readded
        self.document.add_heading(
            'Test: add html with tables, but skip adding tables',
            level=1
        )
        self.parser.options['tables'] = False
        self.parser.add_html_to_document(self.table_html, self.document)

    def test_wrong_argument_type_raises_error(self):
        try:
            self.parser.add_html_to_document(self.document, self.text1)
        except Exception as e:
            assert isinstance(e, ValueError)
            assert "First argument needs to be a <class 'str'>" in str(e)
        else:
            assert False, "Error not raised as expected"

        try:
            self.parser.add_html_to_document(self.text1, self.text1)
        except Exception as e:
            assert isinstance(e, ValueError)
            assert "Second argument" in str(e)
            assert "<class 'docx.document.Document'>" in str(e)
        else:
            assert False, "Error not raised as expected"

    def test_add_html_to_cells_method(self):
        self.document.add_heading(
            'Test: add_html_to_cells method',
            level=1
        )
        table = self.document.add_table(2, 3, style='Table Grid')
        cell = table.cell(0, 0)
        html = '''Line 0 without p tags<p>Line 1 with P tags</p>'''
        self.parser.add_html_to_cell(html, cell)

        cell = table.cell(0, 1)
        html = '''<p>Line 0 with p tags</p>Line 1 without p tags'''
        self.parser.add_html_to_cell(html, cell)

        cell = table.cell(0, 2)
        cell.text = "Pre-defined text that shouldn't be removed."
        html = '''<p>Add HTML to non-empty cell.</p>'''
        self.parser.add_html_to_cell(html, cell)

    def test_inline_code(self):
        self.document.add_heading(
            'Test: inline code block',
            level=1
        )

        html = "<p>This is a sentence that contains <code>some code elements</code> that " \
               "should appear as code.</p>"
        self.parser.add_html_to_document(html, self.document)

    def test_code_block(self):
        self.document.add_heading(
            'Test: code block',
            level=1
        )

        html = """<p><code>
This is a code block.
  That should be NOT be pre-formatted.
It should NOT retain carriage returns,

or blank lines.
</code></p>"""
        self.parser.add_html_to_document(html, self.document)

    def test_pre_block(self):
        self.document.add_heading(
            'Test: pre block',
            level=1
        )

        html = """<pre>
This is a pre-formatted block.
  That should be pre-formatted.
Retaining any carriage returns,

and blank lines.
</pre>
"""
        self.parser.add_html_to_document(html, self.document)

    def test_handling_hr(self):
        self.document.add_heading(
            'Test: Handling of hr',
            level=1
        )
        self.parser.add_html_to_document("<p>paragraph</p><hr><p>paragraph</p>", self.document)

    def test_image_sizes(self):
        self.document.add_heading(
            'Test: Handling of Img sizes',
            level=1
        )
        self.document.add_paragraph("No width")
        self.parser.add_html_to_document(
            "<img src='https://raw.githubusercontent.com/pqzx/h2d/master/testimg.png' />", self.document)

        self.document.add_paragraph("400px")
        self.parser.add_html_to_document(
            "<img src='https://raw.githubusercontent.com/pqzx/h2d/master/testimg.png' width='400' />",
            self.document)

        self.document.add_paragraph("800px, larger than page")
        self.parser.add_html_to_document(
            "<img src='https://raw.githubusercontent.com/pqzx/h2d/master/testimg.png' width='800' />",
            self.document)


        self.document.add_paragraph("An image larger than the page")
        self.parser.add_html_to_document(
            # Attribution: https://commons.wikimedia.org/wiki/Category:Commons_featured_desktop_backgrounds#/media/File:A_storm_at_Pors-Loubous.jpg
            "<img src='https://upload.wikimedia.org/wikipedia/commons/8/88/A_storm_at_Pors-Loubous.jpg' />",
            self.document)

    def test_image_no_src(self):
        self.document.add_heading(
            'Test: Handling IMG without SRC',
            level=1
        )
        self.parser.add_html_to_document("<img />", self.document)

    def test_br_in_table(self):
        # A <br /> in a <table>, but not in a <td>, is illegal.
        self.document.add_heading(
            'Test: Handling BR in TABLE',
            level=1
        )
        self.parser.add_html_to_document(
            "<table><tr>"
            "<td><p>Hello</p></td>"
            "<td><br /><p>Hello</p></td>"
            "<td><p>Hello<br />goodbye</p></td>"
            "</tr></table>",
            self.document)

    def test_leading_br(self):
        self.document.add_heading('Test: Leading BR', level=1)
        self.parser.add_html_to_document("<br /><p>Hello</p>", self.document)


if __name__ == '__main__':
    unittest.main()
