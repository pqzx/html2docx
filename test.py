import unittest
from docx import Document
from h2d import HtmlToDocx

class OutputTest(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        cls.document = Document()
        # cls.parser = HtmlToDocx()
        with open('text1.html', 'r') as tb:
            cls.text1 = tb.read()
        with open('table.html', 'r') as tb:
            cls.table_html = tb.read()

    @classmethod
    def tearDownClass(cls):
        cls.document.save('test.docx')

    def setUp(self):
        self.parser = HtmlToDocx()

    def test_html_with_images_links_style(self):
        self.document.add_heading(
            'Test: add regular html with images, links and some formatting to document',
            level=1
        )
        self.parser.add_html_to_document(self.text1, self.document)

    def test_add_html_to_table_cell(self):
        self.document.add_heading(
            'Test: regular html with images, links, some formatting to table cell',
            level=1
        )
        table = self.document.add_table(2,2, style='Table Grid')
        cell = table.cell(1,1)
        self.parser.add_html_to_document(self.text1, cell)

    def test_add_html_skip_images(self):
        self.document.add_heading(
            'Test: regular html with images, but skip adding images',
            level=1
        )
        self.parser.options['images'] = False
        self.parser.add_html_to_document(self.text1, self.document)

    def test_add_html_with_tables(self):
        self.document.add_heading(
            'Test: add html with tables',
            level=1
        )
        self.parser.add_html_to_document(self.table_html, self.document)

    def test_add_html_skip_tables(self):
        # broken until feature readded
        self.document.add_heading(
            'Test: add html with tables, but skip adding tables',
            level=1
        )
        self.parser.options['tables'] = False
        self.parser.add_html_to_document(self.table_html, self.document)

    def test_wrong_argument_type_raises_error(self):
        try:
            self.parser.add_html_to_document(self.document, self.text1)
        except Exception as e:
            assert isinstance(e, ValueError)
            assert "First argument needs to be a <class 'str'>" in str(e)

        try:
            self.parser.add_html_to_document(self.text1, self.text1)
        except Exception as e:
            assert isinstance(e, ValueError)
            assert "Second argument" in str(e)
            assert "<class 'docx.document.Document'>" in str(e)

    def test_add_html_to_cells_method(self):
        self.document.add_heading(
            'Test: add_html_to_cells method',
            level=1
        )
        table = self.document.add_table(2,2, style='Table Grid')
        cell = table.cell(0,0)
        html = '''Line 0 without p tags<p>Line 1 with P tags</p>'''
        self.parser.add_html_to_cell(html, cell)

        cell = table.cell(0,1)
        html = '''<p>Line 0 with p tags</p>Line 1 without p tags'''
        self.parser.add_html_to_cell(html, cell)

if __name__ == '__main__':
    unittest.main()