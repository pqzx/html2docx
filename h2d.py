"""
Make 'span' in tags dict a stack
maybe do the same for all tags in case of unclosed tags?
optionally use bs4 to clean up invalid html?

the idea is that there is a method that converts html files into docx
but also have api methods that let user have more control e.g. so they
can nest calls to something like 'convert_chunk' in loops

user can pass existing document object as arg 
(if they want to manage rest of document themselves)

How to deal with block level style applied over table elements? e.g. text align
"""
import re, argparse
import io, os
import urllib.request
from urllib.parse import urlparse
from html.parser import HTMLParser

import docx, docx.table
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_COLOR, WD_ALIGN_PARAGRAPH

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

# values in inches
INDENT = 0.25
LIST_INDENT = 0.5
MAX_INDENT = 5.5 # To stop indents going off the page

def get_filename_from_url(url):
    return os.path.basename(urlparse(url).path)

def is_url(url):
    """
    Not to be used for actually validating a url, but in our use case we only 
    care if it's a url or a file path, and they're pretty distinguishable
    """
    parts = urlparse(url)
    return all([parts.scheme, parts.netloc, parts.path])

def fetch_image(url):
    """
    Attempts to fetch an image from a url. 
    If successful returns a bytes object, else returns None

    :return:
    """
    try:
        with urllib.request.urlopen(url) as response:
            # security flaw?
            return io.BytesIO(response.read())
    except urllib.error.URLError:
        return None

def remove_last_occurence(ls, x):
    ls.pop(len(ls) - ls[::-1].index(x) - 1)

def remove_whitespace(string):
    string = re.sub(r'\s*\n\s*', ' ', string)
    return re.sub(r'>\s{2+}<', '><', string)

fonts = {
    'b': 'bold',
    'strong': 'bold',
    'em': 'italic',
    'i': 'italic',
    'u': 'underline',
    's': 'strike',
    'sup': 'superscript',
    'sub': 'subscript',
    'th': 'bold',
}

class HtmlToDocx(HTMLParser):

    def __init__(self):
        super().__init__()
        self.options = {
            'fix-html': True,
            'images': True,
            'tables': True,
            'styles': True,
        }

    def set_initial_attrs(self, document=None):
        self.tags = {
            'span': [],
            'list': [],
        }
        if document:
            self.doc = document
        else:
            self.doc = Document()
        self.bs = self.options['fix-html'] # whether or not to clean with BeautifulSoup
        self.document = self.doc
        self.include_tables = self.options['tables']
        self.include_images = self.options['images']
        self.include_styles = self.options['styles']
        self.paragraph = None
        self.skip = False
        self.skip_tag = None

    def skip_tables_hack(self):
        # temporary hack to deal with nested tables in skipped tables
        if not self.include_tables:
            self.skip = True
            self.skip_tag = 'table'
            return True
        return False

    def add_styles_to_paragraph(self, style):
        if 'text-align' in style:
            align = style['text-align']
            if align == 'center':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif align == 'justify':
                self.paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if 'margin-left' in style:
            margin = style['margin-left']
            units = re.sub(r'[0-9]+', '', margin)
            margin = int(re.sub(r'[a-z]+', '', margin))
            if units == 'px':
                self.paragraph.paragraph_format.left_indent = Inches(min(margin // 10 * INDENT, MAX_INDENT))
            # TODO handle non px units

    def add_styles_to_run(self, style):
        if 'color' in style:
            color = re.sub(r'[a-z()]+', '', style['color'])
            colors = [int(x) for x in color.split(',')]
            self.run.font.color.rgb = RGBColor(*colors)
        if 'background-color' in style:
            color = color = re.sub(r'[a-z()]+', '', style['background-color'])
            colors = [int(x) for x in color.split(',')]
            self.run.font.highlight_color = WD_COLOR.GRAY_25 #TODO: map colors

    def parse_dict_string(self, string, separator=';'):
        new_string = string.replace(" ", '').split(separator)
        string_dict = dict([x.split(':') for x in new_string if ':' in x])
        return string_dict

    def handle_starttag(self, tag, attrs):
        if self.skip:
            return
        current_attrs = dict(attrs)

        if tag == 'span':
            self.tags['span'].append(current_attrs)
            return
        elif tag == 'ol' or tag == 'ul':
            self.tags['list'].append(tag)
            return # don't apply styles for now
        elif tag == 'br':
            self.run.add_break()
            return

        self.tags[tag] = current_attrs
        if tag == 'p':
            self.paragraph = self.doc.add_paragraph()
                        
        elif tag == 'li':
            # check list stack to determine style and depth
            list_depth = len(self.tags['list'])
            if list_depth:
                list_type = self.tags['list'][-1]
            else:
                list_type = 'ul' # assign unordered if no tag

            if list_type == 'ol':
                list_style = "List Number"
            else:
                list_style = 'List Bullet'

            self.paragraph = self.doc.add_paragraph(style=list_style)            
            self.paragraph.paragraph_format.left_indent = Inches(min(list_depth * LIST_INDENT, MAX_INDENT))
            self.paragraph.paragraph_format.line_spacing = 1
            
        elif tag[0] == 'h' and len(tag) == 2:
            if isinstance(self.doc, docx.document.Document):
                h_size = int(tag[1])
                self.paragraph = self.doc.add_heading(level=min(h_size, 9))
            else:
                self.paragraph = self.doc.add_paragraph()

        elif tag == 'img':
            if not self.include_images:
                self.skip = True
                self.skip_tag = tag
                return
            if not isinstance(self.doc, docx.document.Document):
                self.doc.add_paragraph('<image: %s>' % current_attrs['src'])
                return
            src = current_attrs['src']
            # fetch image
            src_is_url = is_url(src)
            if src_is_url:
                try:
                    image = fetch_image(src)
                except urllib.error.URLError:
                    image = None
            else:
                image = src
            # add image to doc
            if image:
                try:
                    self.doc.add_picture(image)
                except FileNotFoundError:
                    image = None
            if not image:
                if src_is_url:
                    self.doc.add_paragraph("<image: %s>" % src)
                else:
                    # avoid exposing filepaths in document
                    self.doc.add_paragraph("<image: %s>" % get_filename_from_url(src))
            # add styles?
            return
        
        elif tag == 'table':
            # TODO: handle nested tables, but for now, just skip
            if isinstance(self.doc, docx.table._Cell):
                self.skip = True
                self.skip_tag = tag
                return
            if not self.include_tables:
                self.skip = True
                self.skip_tag = tag
                return
            # create table with dimensions at current element in self.tables
            rows, cols = self.tables[self.table_no]
            self.table = self.doc.add_table(rows, cols)
            self.table.style = 'Table Grid'
            self.cell_position = [0, 0]
            self.table_styles = current_attrs
            self.paragraph = None

        elif tag == 'tr':
            if self.skip_tables_hack(): return
            self.row_styles = current_attrs
            self.paragraph = None

        elif tag == 'td' or tag == 'th':
            if self.skip_tables_hack(): return
            # point doc to table cell
            self.doc = self.table.cell(*self.cell_position)
            self.paragraph = self.doc.paragraphs[0]
            self.cell_styles = current_attrs
        
        # set new run reference point in case of leading line breaks
        if tag == 'p' or tag == 'li':
            self.run = self.paragraph.add_run()
        
        # add style
        if not self.include_styles:
            return
        if 'style' in current_attrs and self.paragraph:
            style = self.parse_dict_string(current_attrs['style'])
            self.add_styles_to_paragraph(style)

    def handle_endtag(self, tag):
        if self.skip:
            if tag == self.skip_tag:
                self.skip = False
                self.skip_tag = None
                self.paragraph = None
            return
            
        if tag == 'span':
            if self.tags['span']:
                self.tags['span'].pop()
                return
        elif tag == 'ol' or tag == 'ul':
            remove_last_occurence(self.tags['list'], tag)
            return
        elif tag == 'a':
            link = self.tags.pop(tag)
            href = link['href']
            self.paragraph.add_run('<link: %s>' % href)
            return
        elif tag == 'table':
            if self.skip_tables_hack(): return
            self.table_no += 1
            self.table = None
            self.doc = self.document
            self.paragraph = None
            # apply table style across all children
            # actually should probably collate all the styles and apply them in one go
            # at this point, since child styles should override parent styles
            # but here we would do the reverse
        elif tag == 'tr':
            if self.skip_tables_hack(): return
            self.cell_position[0] += 1
            self.cell_position[1] = 0
            # apply row style across all children
        elif tag == 'td' or tag == 'th':
            if self.skip_tables_hack(): return
            self.cell_position[1] += 1
            # apply cell style across all children

        if tag in self.tags:
            self.tags.pop(tag)
        # maybe set relevant reference to None?

    def handle_data(self, data):
        if self.skip:
            return

        if not self.paragraph:
            self.paragraph = self.doc.add_paragraph()

        self.run = self.paragraph.add_run(data)
        spans = self.tags['span']
        for span in spans:
            if 'style' in span:
                style = self.parse_dict_string(span['style'])
                self.add_styles_to_run(style)
        
        # add font style
        for tag in self.tags:
            if tag in fonts:
                font_style = fonts[tag]
                setattr(self.run.font, font_style, True)
 
    def get_tables(self):
        if not hasattr(self, 'soup'):
            self.include_tables = False
            return
            # find other way to do it, or require this dependency?
        tables = self.soup.find_all(['table'])
        self.tables = []
        # get structure of each table. FIXME nested tables also get picked up
        for table in tables:
            rows = table.find_all(['tr'])
            cols = rows[0].find_all(['th', 'td'])
            self.tables.append((len(rows), len(cols)))
        
        self.table_no = 0

    def run_process(self, html):
        if self.bs and BeautifulSoup:
            self.soup = BeautifulSoup(html, 'html.parser')
            html = remove_whitespace(str(self.soup))
        else:
            html = remove_whitespace(html)
        if self.include_tables:
            self.get_tables()
        self.feed(html)

    def add_html_to_document(self, html, document):
        if not isinstance(html, str):
            raise ValueError('First argument needs to be a %s' % str)
        elif not isinstance(document, docx.document.Document) and not isinstance(document, docx.table._Cell):
            raise ValueError('Second argument needs to be a %s' % docx.document.Document)
        self.set_initial_attrs(document)
        self.run_process(html)

    def parse_html_file(self, filename_html, filename_docx=None):
        with open(filename_html, 'r') as infile:
            html = infile.read()
        self.set_initial_attrs()
        self.run_process(html)
        if not filename_docx:
            filename_docx = 'new_docx_file_%s' % filename_html
        self.doc.save('%s.docx' % filename_docx)

if __name__=='__main__':
    
    arg_parser = argparse.ArgumentParser(description='Convert .html file into .docx file with formatting')
    arg_parser.add_argument('filename_html', help='The .html file to be parsed')
    arg_parser.add_argument(
        'filename_docx', 
        nargs='?', 
        help='The name of the .docx file to be saved. Default new_docx_file_[filename_html]', 
        default=None
    )
    arg_parser.add_argument('--bs', action='store_true', 
        help='Attempt to fix html before parsing. Requires bs4. Default True')

    args = vars(arg_parser.parse_args())
    file_html = args.pop('filename_html')
    html_parser = HtmlToDocx()
    html_parser.parse_html_file(file_html, **args)
